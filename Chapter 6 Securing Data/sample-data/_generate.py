"""Generate Wildpaws sample data files for the Chapter 6 Purview demo.

All names, addresses, invoice IDs, and payment numbers are fictional and use
documented test/reserved values so they trigger Purview sensitive info type
(SIT) detectors without exposing real PII.

Test card numbers: 4111 1111 1111 1111 (Visa test), 5500 0000 0000 0004
(Mastercard test), 3782 822463 10005 (Amex test).
Test ABA routing: 021000021, 011000015 (both Federal Reserve public test
routing numbers).
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

OUT_DIR = Path(__file__).parent

CONFIDENTIAL_BANNER = (
    "CONFIDENTIAL - Wildpaws Expeditions Internal Use Only. "
    "Contains customer payment information. Do not share externally."
)
GENERAL_BANNER = (
    "General - Wildpaws Expeditions Internal Reference. Safe for broad "
    "internal sharing and AI agent grounding."
)


def _add_banner(doc: Document, text: str, red: bool) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    if red:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        run.font.color.rgb = RGBColor(0x0B, 0x7A, 0x75)


def _add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _add_p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build_vip_roster() -> Path:
    doc = Document()
    _add_banner(doc, CONFIDENTIAL_BANNER, red=True)
    _add_h1(doc, "Wildpaws Expeditions - VIP Client Roster 2026")
    _add_p(
        doc,
        "This roster tracks Wildpaws VIP clients, their travel companions, "
        "and the payment methods on file for automatic trip deposits. Access "
        "is restricted to the Wildpaws concierge desk. All entries below are "
        "fictional test data.",
    )

    clients = [
        {
            "name": "Priya Nair",
            "member_id": "WPX-VIP-0142",
            "home": "482 Cedar Ridge Rd, Boulder, CO 80302, USA",
            "phone": "+1 (720) 555-0142",
            "email": "priya.nair@example.com",
            "pet": "Kesari, golden retriever, 4 yrs",
            "passport": "N123456789 (USA, expires 2029-08-14)",
            "pet_passport": "PP-KES-0042 (rabies current 2027-03-01)",
            "card_brand": "Visa",
            "card_number": "4111 1111 1111 1111",
            "card_exp": "09/2028",
            "card_cvv": "321",
            "billing_zip": "80302",
        },
        {
            "name": "Diego Fernandez",
            "member_id": "WPX-VIP-0187",
            "home": "17 Rue des Alpes, Chamonix 74400, France",
            "phone": "+33 4 50 55 01 87",
            "email": "diego.fernandez@example.com",
            "pet": "Nube, border collie, 6 yrs",
            "passport": "FR-PA-77218831 (France, expires 2030-01-05)",
            "pet_passport": "PP-NUB-0198 (rabies current 2027-06-11)",
            "card_brand": "Mastercard",
            "card_number": "5500 0000 0000 0004",
            "card_exp": "12/2027",
            "card_cvv": "884",
            "billing_zip": "74400",
        },
        {
            "name": "Amaka Obi",
            "member_id": "WPX-VIP-0203",
            "home": "9 Kingsway Cres, Toronto ON M6P 2K1, Canada",
            "phone": "+1 (416) 555-0203",
            "email": "amaka.obi@example.com",
            "pet": "Zuri, rhodesian ridgeback, 3 yrs",
            "passport": "CA-XA987654 (Canada, expires 2028-11-22)",
            "pet_passport": "PP-ZUR-0311 (rabies current 2026-12-04)",
            "card_brand": "American Express",
            "card_number": "3782 822463 10005",
            "card_exp": "05/2029",
            "card_cvv": "9142",
            "billing_zip": "M6P2K1",
        },
    ]

    for c in clients:
        _add_h2(doc, f"{c['name']} ({c['member_id']})")
        _add_p(doc, f"Home address: {c['home']}")
        _add_p(doc, f"Phone: {c['phone']}")
        _add_p(doc, f"Email: {c['email']}")
        _add_p(doc, f"Travel companion: {c['pet']}")
        _add_p(doc, f"Passport: {c['passport']}")
        _add_p(doc, f"Pet passport: {c['pet_passport']}")
        _add_p(
            doc,
            f"Payment on file: {c['card_brand']} {c['card_number']}, "
            f"exp {c['card_exp']}, CVV {c['card_cvv']}, "
            f"billing ZIP {c['billing_zip']}.",
        )

    out = OUT_DIR / "Wildpaws_VIP_Client_Roster_2026.docx"
    doc.save(out)
    return out


def build_vendor_invoice() -> Path:
    doc = Document()
    _add_banner(doc, CONFIDENTIAL_BANNER, red=True)
    _add_h1(doc, "Vendor Invoice - Banff Basecamp Lodge")
    _add_p(doc, "Invoice ID: WPX-INV-2001")
    _add_p(doc, "Vendor: Banff Basecamp Lodge Ltd. (fictional)")
    _add_p(doc, "Bill to: Wildpaws Expeditions Inc.")
    _add_p(doc, "Issued: 2026-07-14 - Due: 2026-08-13")
    _add_p(doc, "Amount due: $12,450.00 USD")

    _add_h2(doc, "Line items")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Nights"
    hdr[2].text = "Total"
    rows = [
        ("Pet-friendly cabin block, guests + dogs", "5", "$8,750.00"),
        ("Guided trail package (Sulphur Mountain)", "2", "$2,400.00"),
        ("Pet gear rental and cleaning fee", "-", "$1,300.00"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v

    _add_h2(doc, "Payment instructions (TEST - Not Real)")
    _add_p(
        doc,
        "Please wire the balance to the account below within 30 days:",
    )
    _add_p(doc, "Beneficiary: Banff Basecamp Lodge Ltd.")
    _add_p(doc, "Bank: First Reserve Bank of Alberta (fictional)")
    _add_p(doc, "ABA Routing Number: 021000021")
    _add_p(doc, "Bank Account Number: 000123456789")
    _add_p(doc, "SWIFT: FRSVCAT2XXX")
    _add_p(
        doc,
        "Backup credit card on file (chargeback only): Visa "
        "4111 1111 1111 1111, exp 09/2028.",
    )

    out = OUT_DIR / "Wildpaws_Vendor_Invoice_BanffLodge_INV-2001.docx"
    doc.save(out)
    return out


def build_trip_deposits_ledger() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Trip Deposits"

    banner = "CONFIDENTIAL - Wildpaws Expeditions payment records. Test data."
    ws["A1"] = banner
    ws["A1"].font = Font(bold=True, color="C00000")
    ws.merge_cells("A1:H1")

    headers = [
        "Guest",
        "Member ID",
        "Trip",
        "Departure",
        "Card Brand",
        "Card Number",
        "ABA Routing",
        "Deposit (USD)",
    ]
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=i, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(
            start_color="0B7A75", end_color="0B7A75", fill_type="solid"
        )
        cell.alignment = Alignment(horizontal="center")

    rows = [
        (
            "Priya Nair",
            "WPX-VIP-0142",
            "Banff Autumn Expedition",
            "2026-09-18",
            "Visa",
            "4111 1111 1111 1111",
            "021000021",
            2450.00,
        ),
        (
            "Diego Fernandez",
            "WPX-VIP-0187",
            "Chamonix Winter Trek",
            "2026-12-02",
            "Mastercard",
            "5500 0000 0000 0004",
            "011000015",
            3175.50,
        ),
        (
            "Amaka Obi",
            "WPX-VIP-0203",
            "Patagonia Long Trail",
            "2027-02-11",
            "American Express",
            "3782 822463 10005",
            "021000021",
            4980.00,
        ),
        (
            "Marta Kowalski",
            "WPX-STD-1188",
            "Tokyo Neon and Parks",
            "2026-10-05",
            "Visa",
            "4111 1111 1111 1111",
            "011000015",
            1600.00,
        ),
        (
            "Jerome Blake",
            "WPX-STD-1201",
            "Banff Autumn Expedition",
            "2026-09-18",
            "Mastercard",
            "5500 0000 0000 0004",
            "021000021",
            2450.00,
        ),
    ]
    for r_idx, row in enumerate(rows, start=4):
        for c_idx, val in enumerate(row, start=1):
            ws.cell(row=r_idx, column=c_idx, value=val)

    for col_letter, width in zip(
        ["A", "B", "C", "D", "E", "F", "G", "H"],
        [22, 16, 28, 12, 14, 22, 14, 14],
    ):
        ws.column_dimensions[col_letter].width = width

    out = OUT_DIR / "Wildpaws_Trip_Deposits_Ledger.xlsx"
    wb.save(out)
    return out


def build_public_trail_catalog() -> Path:
    doc = Document()
    _add_banner(doc, GENERAL_BANNER, red=False)
    _add_h1(doc, "Wildpaws Public Trail Catalog")
    _add_p(
        doc,
        "This catalog lists Wildpaws-endorsed pet-friendly trails cleared "
        "for use in marketing content and AI agent grounding. It contains no "
        "customer or payment information.",
    )

    trails = [
        {
            "name": "Sulphur Mountain Boardwalk",
            "region": "Banff, Alberta, Canada",
            "difficulty": "Easy",
            "length": "5.5 km round trip",
            "season": "May to October",
            "pets": "Dogs on leash welcome; water bowls at summit lodge",
            "notes": "Gondola return option, wheelchair-accessible upper deck",
        },
        {
            "name": "Aiguille du Midi Panoramic Loop",
            "region": "Chamonix, France",
            "difficulty": "Moderate",
            "length": "8 km loop",
            "season": "June to September",
            "pets": (
                "Dogs allowed on lower sections; carriers required in "
                "cable cars"
            ),
            "notes": "Altitude 2300 m, bring water and pet booties",
        },
        {
            "name": "Nakasendo Old Post Road",
            "region": "Nagano, Japan",
            "difficulty": "Easy to moderate",
            "length": "8 km point-to-point (Magome to Tsumago)",
            "season": "April to November",
            "pets": (
                "Small dogs in carriers permitted; large dogs must remain "
                "on marked side paths"
            ),
            "notes": "Cobblestone and forest, minshuku pet-friendly options",
        },
        {
            "name": "Torres del Paine W Circuit (Day 1 sample)",
            "region": "Patagonia, Chile",
            "difficulty": "Strenuous",
            "length": "18 km round trip",
            "season": "November to March",
            "pets": (
                "Dogs not permitted inside the national park; kennel "
                "referral available at Puerto Natales"
            ),
            "notes": "Guests only; strong winds, dress in layers",
        },
        {
            "name": "Uinta Aspen Loop",
            "region": "Uinta Mountains, Utah, USA",
            "difficulty": "Easy",
            "length": "4 km loop",
            "season": "June to October",
            "pets": "All friendly dogs on leash; horse trailhead nearby",
            "notes": "Wildpaws Basecamp partner trail, ranger station on-site",
        },
    ]

    for t in trails:
        _add_h2(doc, t["name"])
        _add_p(doc, f"Region: {t['region']}")
        _add_p(doc, f"Difficulty: {t['difficulty']}")
        _add_p(doc, f"Length: {t['length']}")
        _add_p(doc, f"Season: {t['season']}")
        _add_p(doc, f"Pet policy: {t['pets']}")
        _add_p(doc, f"Notes: {t['notes']}")

    out = OUT_DIR / "Wildpaws_Public_Trail_Catalog.docx"
    doc.save(out)
    return out


def build_packing_guide() -> Path:
    doc = Document()
    _add_banner(doc, GENERAL_BANNER, red=False)
    _add_h1(doc, "Wildpaws Packing Guide - Traveling with Pets")
    _add_p(
        doc,
        "General guidance Wildpaws concierges share with guests. No PII, "
        "no payment details. Safe for AI agent grounding.",
    )

    _add_h2(doc, "Documents to carry")
    for item in [
        "Owner passport and valid travel visa where required",
        "Pet passport with current rabies certificate",
        "Vet health certificate issued within 10 days of departure",
        "Wildpaws trip itinerary and lodge confirmations",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_h2(doc, "Gear for cold-weather expeditions")
    for item in [
        "Insulated dog jacket sized to chest circumference",
        "Paw booties with grip soles for icy trails",
        "Collapsible bowl and 1 L insulated water bottle",
        "High-visibility LED collar and spare leash",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_h2(doc, "Gear for warm-weather expeditions")
    for item in [
        "Cooling vest and breathable harness",
        "Portable shade tarp and stakes",
        "Electrolyte pet drink packets",
        "Tick removal kit and pet-safe repellent",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    _add_h2(doc, "Universal health kit")
    for item in [
        "Prescribed medications with dosage schedule",
        "Bandages, styptic powder, and blunt tip scissors",
        "Emergency vet contact card for destination region",
        "Wildpaws 24 hour concierge card",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    out = OUT_DIR / "Wildpaws_Packing_Guide_Pets.docx"
    doc.save(out)
    return out


def build_employee_expenses() -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "Trail Expenses"

    banner = (
        "CONFIDENTIAL - Wildpaws employee reimbursement ledger. Test data."
    )
    ws["A1"] = banner
    ws["A1"].font = Font(bold=True, color="C00000")
    ws.merge_cells("A1:G1")

    headers = [
        "Employee",
        "Employee ID",
        "Trip",
        "Date",
        "Card Number",
        "Amount (USD)",
        "Reimbursement Status",
    ]
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=3, column=i, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(
            start_color="0B7A75", end_color="0B7A75", fill_type="solid"
        )
        cell.alignment = Alignment(horizontal="center")

    rows = [
        (
            "Sam Iverson",
            "WPX-EMP-0071",
            "Banff Basecamp Setup",
            "2026-07-12",
            "4111 1111 1111 1111",
            842.15,
            "Approved",
        ),
        (
            "Linh Tran",
            "WPX-EMP-0089",
            "Chamonix Pre-Season Scout",
            "2026-07-19",
            "5500 0000 0000 0004",
            1204.60,
            "Pending",
        ),
        (
            "Bode Ellison",
            "WPX-EMP-0104",
            "Uinta Ranger Sync",
            "2026-07-22",
            "3782 822463 10005",
            356.90,
            "Approved",
        ),
        (
            "Nia Okafor",
            "WPX-EMP-0118",
            "Patagonia Route Survey",
            "2026-07-25",
            "4111 1111 1111 1111",
            2088.00,
            "Pending",
        ),
    ]
    for r_idx, row in enumerate(rows, start=4):
        for c_idx, val in enumerate(row, start=1):
            ws.cell(row=r_idx, column=c_idx, value=val)

    for col_letter, width in zip(
        ["A", "B", "C", "D", "E", "F", "G"],
        [18, 16, 30, 12, 22, 14, 22],
    ):
        ws.column_dimensions[col_letter].width = width

    out = OUT_DIR / "Wildpaws_Employee_Trail_Expenses.xlsx"
    wb.save(out)
    return out


def main() -> None:
    created = [
        build_vip_roster(),
        build_vendor_invoice(),
        build_trip_deposits_ledger(),
        build_public_trail_catalog(),
        build_packing_guide(),
        build_employee_expenses(),
    ]
    print("Created files:")
    for p in created:
        size_kb = os.path.getsize(p) / 1024
        print(f"  {p.name}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
